"""
Fill Application - DOCX Generator Service

Generates DOCX documents from data rows with flexible formatting options.
Supports creating documents from scratch with tables, paragraphs, and styling.
"""

from io import BytesIO
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxGeneratorError(Exception):
    """
    Custom exception for DOCX generation errors.
    """

    def __init__(self, message: str) -> None:
        """
        Initialize DOCX generator error.

        Args:
            message: Error message describing what went wrong
        """
        self.message = message
        super().__init__(self.message)


class DocxGenerator:
    """
    Generates DOCX documents from data rows.

    Features:
    - Create documents with tables, paragraphs, and headings
    - Support for various data types (strings, numbers, dates, booleans)
    - Flexible formatting options (font size, alignment, colors)
    - Output to file or return as bytes
    - Template-based generation with placeholders
    """

    def __init__(self) -> None:
        """
        Initialize DOCX generator.

        Raises:
            ImportError: If python-docx is not available
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not available. "
                "Install it with: pip install python-docx"
            )

    def generate_from_data(
        self,
        data: list[dict[str, Any]],
        output_path: str | Path | None = None,
        title: str | None = None,
        include_headers: bool = True,
    ) -> bytes | None:
        """
        Generate a DOCX document from a list of data rows.

        Creates a document with a title (optional) and a table containing
        all data rows. Headers are automatically generated from column names.

        Args:
            data: List of dictionaries with consistent column names
            output_path: Optional path to save the document
            title: Optional document title/heading
            include_headers: Whether to include table headers (default: True)

        Returns:
            Generated document as bytes (or None if output_path provided)

        Raises:
            DocxGeneratorError: If data is empty or invalid
            ValueError: If output path is invalid
        """
        if not data:
            raise DocxGeneratorError("Data cannot be empty")

        if not isinstance(data, list):
            raise DocxGeneratorError(
                f"Data must be a list, got {type(data)}"
            )

        if not all(isinstance(row, dict) for row in data):
            raise DocxGeneratorError(
                "All data rows must be dictionaries"
            )

        # Create new document
        doc = Document()

        # Add title if provided
        if title:
            self._add_title(doc, title)

        # Get all column names from data
        columns = self._get_columns_from_data(data)

        # Add table with data
        self._add_data_table(
            doc, data, columns, include_headers
        )

        # Save or return bytes
        return self._save_document(doc, output_path)

    def generate_from_template(
        self,
        template_path: str | Path,
        data: dict[str, Any],
        output_path: str | Path | None = None,
    ) -> bytes | None:
        """
        Generate a DOCX document from a template file.

        Fills placeholders in a template document with data values.
        Placeholder syntax: {{field_name}}

        Args:
            template_path: Path to template DOCX file
            data: Dictionary of field names to values
            output_path: Optional path to save the document

        Returns:
            Generated document as bytes (or None if output_path provided)

        Raises:
            DocxGeneratorError: If template is invalid
            FileNotFoundError: If template file doesn't exist
        """
        template_path = Path(template_path)

        if not template_path.exists():
            raise FileNotFoundError(
                f"Template file not found: {template_path}"
            )

        if template_path.suffix.lower() != ".docx":
            raise DocxGeneratorError(
                f"Template must be .docx file, got {template_path.suffix}"
            )

        # Load template
        try:
            doc = Document(str(template_path))
        except Exception as e:
            raise DocxGeneratorError(
                f"Failed to load template: {e}"
            )

        # Fill placeholders
        self._fill_placeholders(doc, data)

        # Save or return bytes
        return self._save_document(doc, output_path)

    def generate_simple_document(
        self,
        content: str | list[str],
        output_path: str | Path | None = None,
        title: str | None = None,
    ) -> bytes | None:
        """
        Generate a simple DOCX document with text content.

        Creates a document with plain text paragraphs.
        Content can be a single string or a list of strings (one per paragraph).

        Args:
            content: Text content (string or list of strings)
            output_path: Optional path to save the document
            title: Optional document title/heading

        Returns:
            Generated document as bytes (or None if output_path provided)

        Raises:
            DocxGeneratorError: If content is invalid
        """
        if not content:
            raise DocxGeneratorError("Content cannot be empty")

        # Create new document
        doc = Document()

        # Add title if provided
        if title:
            self._add_title(doc, title)

        # Add content paragraphs
        if isinstance(content, str):
            # Single paragraph
            doc.add_paragraph(content)
        elif isinstance(content, list):
            # Multiple paragraphs
            for paragraph_text in content:
                if paragraph_text:  # Skip empty strings
                    doc.add_paragraph(paragraph_text)
        else:
            raise DocxGeneratorError(
                f"Content must be string or list, got {type(content)}"
            )

        # Save or return bytes
        return self._save_document(doc, output_path)

    def _add_title(self, doc: Document, title: str) -> None:
        """
        Add a title heading to the document.

        Args:
            doc: Document object
            title: Title text
        """
        heading = doc.add_heading(title, level=1)
        # Apply some basic styling
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True

    def _get_columns_from_data(
        self, data: list[dict[str, Any]]
    ) -> list[str]:
        """
        Extract all column names from data rows.

        Args:
            data: List of data rows

        Returns:
            List of column names
        """
        columns = set()
        for row in data:
            columns.update(row.keys())
        return sorted(list(columns))

    def _add_data_table(
        self,
        doc: Document,
        data: list[dict[str, Any]],
        columns: list[str],
        include_headers: bool,
    ) -> None:
        """
        Add a table with data to the document.

        Args:
            doc: Document object
            data: List of data rows
            columns: List of column names
            include_headers: Whether to include header row
        """
        if not data:
            return

        # Calculate table size
        num_rows = len(data) + (1 if include_headers else 0)
        num_cols = len(columns)

        # Add table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        # Add headers if requested
        if include_headers:
            header_cells = table.rows[0].cells
            for i, column in enumerate(columns):
                header_cells[i].text = str(column)
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

        # Add data rows
        data_start_row = 1 if include_headers else 0
        for i, row_data in enumerate(data):
            row_cells = table.rows[data_start_row + i].cells
            for j, column in enumerate(columns):
                value = row_data.get(column, "")
                # Convert value to string (handle None, numbers, etc.)
                if value is None:
                    value = ""
                row_cells[j].text = str(value)

    def _fill_placeholders(
        self, doc: Document, data: dict[str, Any]
    ) -> None:
        """
        Fill placeholders in document with data values.

        Args:
            doc: Document object
            data: Dictionary of field names to values
        """
        import re

        # Placeholder pattern: {{field_name}}
        pattern = re.compile(r"\{\{([a-zA-Z0-9_\-\s]+?)\}\}")

        def replace_placeholder(match):
            """Replace placeholder with data value."""
            field_name = match.group(1).strip()
            value = data.get(field_name, match.group(0))  # Keep if not found
            return str(value) if value is not None else ""

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = pattern.sub(replace_placeholder, run.text)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = pattern.sub(
                                    replace_placeholder, run.text
                                )

    def _save_document(
        self, doc: Document, output_path: str | Path | None
    ) -> bytes | None:
        """
        Save document to file or return as bytes.

        Args:
            doc: Document object
            output_path: Optional path to save document

        Returns:
            Document as bytes (or None if output_path provided)
        """
        if output_path:
            output_path = Path(output_path)
            # Create parent directories if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            return None
        else:
            # Return as bytes
            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()


def generate_docx_from_data(
    data: list[dict[str, Any]],
    output_path: str | Path | None = None,
    title: str | None = None,
    include_headers: bool = True,
) -> bytes | None:
    """
    Convenience function to generate a DOCX document from data rows.

    Args:
        data: List of dictionaries with consistent column names
        output_path: Optional path to save the document
        title: Optional document title
        include_headers: Whether to include table headers (default: True)

    Returns:
        Generated document as bytes (or None if output_path provided)

    Raises:
        DocxGeneratorError: If data is empty or invalid
    """
    generator = DocxGenerator()
    return generator.generate_from_data(
        data, output_path, title, include_headers
    )


def generate_docx_from_template(
    template_path: str | Path,
    data: dict[str, Any],
    output_path: str | Path | None = None,
) -> bytes | None:
    """
    Convenience function to generate a DOCX document from a template.

    Args:
        template_path: Path to template DOCX file
        data: Dictionary of field names to values
        output_path: Optional path to save the document

    Returns:
        Generated document as bytes (or None if output_path provided)

    Raises:
        DocxGeneratorError: If template is invalid
        FileNotFoundError: If template file doesn't exist
    """
    generator = DocxGenerator()
    return generator.generate_from_template(
        template_path, data, output_path
    )
