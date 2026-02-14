"""
Fill Application - Placeholder Parser Service

Extracts {{field_name}} patterns from template content.
Supports text and DOCX file parsing for placeholder detection.
"""

import re
from pathlib import Path
from typing import Any

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class PlaceholderParser:
    """
    Parser for extracting {{field_name}} placeholders from templates.

    Placeholders use the double-brace syntax: {{field_name}}
    - Field names can contain: letters, numbers, underscores, hyphens
    - Placeholders can span multiple lines
    - Duplicate placeholders are all returned (use set() for uniqueness)

    Example:
        "Dear {{customer_name}}, your invoice {{invoice_number}} is due."
        -> ["customer_name", "invoice_number"]
    """

    # Regex pattern for {{field_name}} placeholders
    # Matches: {{field_name}}
    # - field_name: alphanumeric, underscores, hyphens
    # - Supports multiline placeholders
    # - Non-greedy matching to handle multiple placeholders
    PATTERN = re.compile(r"\{\{([a-zA-Z0-9_\-\s]+?)\}\}", re.MULTILINE)

    # Pattern for extracting just the field name (trimmed)
    FIELD_NAME_PATTERN = re.compile(r"^[a-zA-Z0-9_\-\s]+$")

    def __init__(self) -> None:
        """
        Initialize the placeholder parser.

        Raises:
            ImportError: If python-docx is not available (for DOCX parsing)
        """
        self._docx_available = DOCX_AVAILABLE

    def extract_from_text(self, content: str) -> list[str]:
        """
        Extract placeholders from text content.

        Args:
            content: Text content containing placeholders

        Returns:
            List of placeholder field names (preserves order, may include duplicates)
            Field names are trimmed of leading/trailing whitespace

        Raises:
            ValueError: If content is None or empty
        """
        if content is None:
            raise ValueError("Content cannot be None")

        if not isinstance(content, str):
            raise ValueError(f"Content must be string, got {type(content)}")

        # Find all placeholders
        matches = self.PATTERN.findall(content)

        # Trim whitespace from each field name
        placeholders = [match.strip() for match in matches]

        return placeholders

    def extract_from_file(self, file_path: str | Path) -> list[str]:
        """
        Extract placeholders from a template file.

        Supports:
        - Text files (.txt, .csv, .md, .html, etc.)
        - DOCX files (.docx) if python-docx is available

        Args:
            file_path: Path to template file

        Returns:
            List of placeholder field names

        Raises:
            ValueError: If file doesn't exist or unsupported format
            ImportError: If parsing DOCX but python-docx not available
        """
        file_path = Path(file_path)

        # Check file exists
        if not file_path.exists():
            raise ValueError(f"File not found: {file_path}")

        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")

        # Determine file type and extract accordingly
        extension = file_path.suffix.lower()

        if extension == ".docx":
            return self._extract_from_docx(file_path)
        else:
            # Treat as text file
            try:
                content = file_path.read_text(encoding="utf-8")
                return self.extract_from_text(content)
            except UnicodeDecodeError:
                raise ValueError(
                    f"Could not read file as text (encoding issue): {file_path}"
                )

    def extract_unique_from_text(self, content: str) -> list[str]:
        """
        Extract unique placeholders from text content.

        Same as extract_from_text() but returns only unique placeholders,
        preserving first occurrence order.

        Args:
            content: Text content containing placeholders

        Returns:
            List of unique placeholder field names (no duplicates)
        """
        placeholders = self.extract_from_text(content)

        # Preserve order while removing duplicates
        seen: set[str] = set()
        unique: list[str] = []
        for placeholder in placeholders:
            if placeholder not in seen:
                seen.add(placeholder)
                unique.append(placeholder)

        return unique

    def extract_unique_from_file(self, file_path: str | Path) -> list[str]:
        """
        Extract unique placeholders from a template file.

        Same as extract_from_file() but returns only unique placeholders,
        preserving first occurrence order.

        Args:
            file_path: Path to template file

        Returns:
            List of unique placeholder field names (no duplicates)
        """
        placeholders = self.extract_from_file(file_path)

        # Preserve order while removing duplicates
        seen: set[str] = set()
        unique: list[str] = []
        for placeholder in placeholders:
            if placeholder not in seen:
                seen.add(placeholder)
                unique.append(placeholder)

        return unique

    def validate_placeholder(self, placeholder: str) -> bool:
        """
        Validate a placeholder field name.

        Args:
            placeholder: Field name to validate

        Returns:
            True if placeholder is valid, False otherwise

        Valid field names:
        - Not empty
        - Contains only: letters, numbers, underscores, hyphens, spaces
        - No double braces {{}}
        """
        if not placeholder or not isinstance(placeholder, str):
            return False

        # Check if it matches the valid pattern
        return bool(self.FIELD_NAME_PATTERN.match(placeholder))

    def _extract_from_docx(self, file_path: Path) -> list[str]:
        """
        Extract placeholders from a DOCX file.

        Reads all paragraphs and tables from the DOCX file,
        then extracts placeholders from the combined text.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of placeholder field names

        Raises:
            ImportError: If python-docx is not available
            ValueError: If file is not a valid DOCX
        """
        if not self._docx_available:
            raise ImportError(
                "python-docx is required to parse DOCX files. "
                "Install with: pip install python-docx"
            )

        try:
            doc = Document(str(file_path))
        except Exception as e:
            raise ValueError(f"Failed to open DOCX file: {e}")

        # Collect all text from paragraphs and tables
        all_text_parts: list[str] = []

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                all_text_parts.append(paragraph.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        all_text_parts.append(cell.text)

        # Combine all text and extract placeholders
        combined_text = "\n".join(all_text_parts)
        return self.extract_from_text(combined_text)


# Global singleton instance
_default_parser: PlaceholderParser | None = None


def get_placeholder_parser() -> PlaceholderParser:
    """
    Get the global placeholder parser singleton.

    Returns:
        The global PlaceholderParser instance (created if needed)
    """
    global _default_parser

    if _default_parser is None:
        _default_parser = PlaceholderParser()

    return _default_parser
