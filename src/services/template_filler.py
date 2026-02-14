"""
Fill Application - Template Filler Service

Fills template placeholders with data values from source files.
Supports text-based templates and DOCX templates.
"""

import re
from pathlib import Path
from typing import Any

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.models.mapping import Mapping


class TemplateFillerError(Exception):
    """
    Custom exception for template filling errors.
    """

    def __init__(self, message: str) -> None:
        """
        Initialize template filler error.

        Args:
            message: Error message describing what went wrong
        """
        self.message = message
        super().__init__(self.message)


class TemplateFiller:
    """
    Fills template placeholders with data values.

    Supports:
    - Text-based templates (.txt, .md, .csv, .html, etc.)
    - DOCX templates (.docx) if python-docx is available

    Placeholder syntax: {{field_name}}
    - Replaces {{field_name}} with corresponding data value
    - Handles missing values (keeps placeholder or uses default)
    - Preserves template formatting
    """

    # Regex pattern for {{field_name}} placeholders
    # Matches: {{field_name}}
    # - field_name: alphanumeric, underscores, hyphens
    # - Supports multiline placeholders
    # - Non-greedy matching to handle multiple placeholders
    PLACEHOLDER_PATTERN = re.compile(r"\{\{([a-zA-Z0-9_\-\s]+?)\}\}", re.MULTILINE)

    def __init__(
        self, missing_placeholder_strategy: str = "keep"
    ) -> None:
        """
        Initialize the template filler.

        Args:
            missing_placeholder_strategy: Strategy for handling missing values
                - "keep": Keep the placeholder in the output (default)
                - "empty": Replace with empty string
                - "default": Replace with a default value (e.g., "N/A")

        Raises:
            ValueError: If missing_placeholder_strategy is invalid
            ImportError: If python-docx is not available (for DOCX filling)
        """
        valid_strategies = ["keep", "empty", "default"]
        if missing_placeholder_strategy not in valid_strategies:
            raise ValueError(
                f"Invalid missing_placeholder_strategy: {missing_placeholder_strategy}. "
                f"Must be one of: {', '.join(valid_strategies)}"
            )

        self._missing_strategy = missing_placeholder_strategy
        self._default_value = "N/A"

    def fill_text_template(
        self,
        template_content: str,
        data_row: dict[str, Any],
        mapping: Mapping,
    ) -> str:
        """
        Fill a text-based template with data values.

        Args:
            template_content: Template content as string with {{placeholders}}
            data_row: Dictionary of column names to data values
            mapping: Mapping object defining column -> placeholder mappings

        Returns:
            Filled template content as string

        Raises:
            TemplateFillerError: If template content is invalid
            ValueError: If data_row is not a dictionary
        """
        if template_content is None:
            raise TemplateFillerError("Template content cannot be None")

        if not isinstance(template_content, str):
            raise TemplateFillerError(
                f"Template content must be string, got {type(template_content)}"
            )

        if not isinstance(data_row, dict):
            raise ValueError(
                f"data_row must be dictionary, got {type(data_row)}"
            )

        # Build placeholder -> value mapping
        placeholder_values = self._build_placeholder_values(
            data_row, mapping
        )

        # Replace placeholders in template
        def replace_placeholder(match: re.Match) -> str:
            """Replace a single placeholder with its value."""
            placeholder_name = match.group(1).strip()

            # Get value for this placeholder
            if placeholder_name in placeholder_values:
                value = placeholder_values[placeholder_name]

                # Convert value to string (handle None, numbers, etc.)
                if value is None:
                    return self._handle_missing_value(placeholder_name)
                return str(value)

            # Placeholder not found in mapping - handle based on strategy
            return self._handle_missing_value(placeholder_name)

        # Replace all placeholders
        filled_content = self.PLACEHOLDER_PATTERN.sub(
            replace_placeholder, template_content
        )

        return filled_content

    def fill_docx_template(
        self,
        template_path: str | Path,
        data_row: dict[str, Any],
        mapping: Mapping,
        output_path: str | Path | None = None,
    ) -> bytes:
        """
        Fill a DOCX template with data values.

        Args:
            template_path: Path to DOCX template file
            data_row: Dictionary of column names to data values
            mapping: Mapping object defining column -> placeholder mappings
            output_path: Optional path to save filled document (if None, returns bytes)

        Returns:
            Filled DOCX document as bytes (or None if output_path provided)

        Raises:
            ImportError: If python-docx is not available
            TemplateFillerError: If template file is invalid or cannot be read
            FileNotFoundError: If template file doesn't exist
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not available. "
                "Install it with: pip install python-docx"
            )

        # Convert path to Path object
        template_path = Path(template_path)

        # Check template file exists
        if not template_path.exists():
            raise FileNotFoundError(
                f"Template file not found: {template_path}"
            )

        if not template_path.is_file():
            raise TemplateFillerError(
                f"Template path is not a file: {template_path}"
            )

        # Load the document
        try:
            doc = Document(str(template_path))
        except Exception as e:
            raise TemplateFillerError(
                f"Failed to load DOCX template: {e}"
            )

        # Build placeholder -> value mapping
        placeholder_values = self._build_placeholder_values(
            data_row, mapping
        )

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                self._replace_placeholders_in_text(run, placeholder_values)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._replace_placeholders_in_text(
                                run, placeholder_values
                            )

        # Save or return bytes
        if output_path:
            output_path = Path(output_path)
            doc.save(str(output_path))
            return None
        else:
            # Return as bytes
            from io import BytesIO

            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

    def fill_template(
        self,
        template_path: str | Path,
        data_row: dict[str, Any],
        mapping: Mapping,
    ) -> bytes:
        """
        Fill a template with data values (auto-detects format).

        Args:
            template_path: Path to template file (text or DOCX)
            data_row: Dictionary of column names to data values
            mapping: Mapping object defining column -> placeholder mappings

        Returns:
            Filled template as bytes

        Raises:
            TemplateFillerError: If template format is not supported
            FileNotFoundError: If template file doesn't exist
        """
        template_path = Path(template_path)

        # Check file exists
        if not template_path.exists():
            raise FileNotFoundError(
                f"Template file not found: {template_path}"
            )

        # Determine format and fill accordingly
        suffix = template_path.suffix.lower()

        if suffix == ".docx":
            return self.fill_docx_template(
                template_path, data_row, mapping
            )
        elif suffix in [".txt", ".md", ".html", ".csv"]:
            # Text-based template
            with open(template_path, "r", encoding="utf-8") as f:
                template_content = f.read()

            filled_content = self.fill_text_template(
                template_content, data_row, mapping
            )

            return filled_content.encode("utf-8")
        else:
            raise TemplateFillerError(
                f"Unsupported template format: {suffix}. "
                f"Supported formats: .txt, .md, .html, .csv, .docx"
            )

    def _build_placeholder_values(
        self,
        data_row: dict[str, Any],
        mapping: Mapping,
    ) -> dict[str, Any]:
        """
        Build placeholder -> value mapping from data row and column mapping.

        Args:
            data_row: Dictionary of column names to data values
            mapping: Mapping object with column -> placeholder mappings

        Returns:
            Dictionary of placeholder names to data values
        """
        placeholder_values = {}

        for column_name, placeholder_name in mapping.column_mappings.items():
            # Get value from data row (handle missing columns)
            value = data_row.get(column_name)
            placeholder_values[placeholder_name] = value

        return placeholder_values

    def _handle_missing_value(self, placeholder_name: str) -> str:
        """
        Handle a missing value for a placeholder.

        Args:
            placeholder_name: Name of the placeholder missing a value

        Returns:
            String to use in place of the missing value
        """
        if self._missing_strategy == "keep":
            # Keep the placeholder in output
            return f"{{{{{placeholder_name}}}}}"
        elif self._missing_strategy == "empty":
            # Replace with empty string
            return ""
        else:  # "default"
            # Replace with default value
            return self._default_value

    def _replace_placeholders_in_text(
        self,
        text_element: Any,
        placeholder_values: dict[str, Any],
    ) -> None:
        """
        Replace placeholders in a text element (for DOCX).

        Args:
            text_element: Text element to replace placeholders in (e.g., docx run)
            placeholder_values: Dictionary of placeholder names to values
        """
        if not hasattr(text_element, "text"):
            return

        original_text = text_element.text
        if not original_text:
            return

        def replace_placeholder(match: re.Match) -> str:
            """Replace a single placeholder with its value."""
            placeholder_name = match.group(1).strip()

            # Get value for this placeholder
            if placeholder_name in placeholder_values:
                value = placeholder_values[placeholder_name]

                # Convert value to string
                if value is None:
                    return self._handle_missing_value(placeholder_name)
                return str(value)

            # Placeholder not found - handle based on strategy
            return self._handle_missing_value(placeholder_name)

        # Replace placeholders
        new_text = self.PLACEHOLDER_PATTERN.sub(
            replace_placeholder, original_text
        )

        text_element.text = new_text


def fill_template(
    template_path: str | Path,
    data_row: dict[str, Any],
    mapping: Mapping,
    missing_placeholder_strategy: str = "keep",
) -> bytes:
    """
    Convenience function to fill a template with data values.

    Args:
        template_path: Path to template file
        data_row: Dictionary of column names to data values
        mapping: Mapping object defining column -> placeholder mappings
        missing_placeholder_strategy: Strategy for handling missing values
            - "keep": Keep the placeholder in the output (default)
            - "empty": Replace with empty string
            - "default": Replace with "N/A"

    Returns:
        Filled template as bytes

    Raises:
        TemplateFillerError: If template cannot be filled
        FileNotFoundError: If template file doesn't exist
    """
    filler = TemplateFiller(
        missing_placeholder_strategy=missing_placeholder_strategy
    )
    return filler.fill_template(template_path, data_row, mapping)
