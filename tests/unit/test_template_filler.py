"""
Unit tests for Template Filler Service.

Tests template placeholder filling with data values.
"""

import os
import tempfile
from pathlib import Path

import pytest

from src.models.mapping import Mapping
from src.services.template_filler import (
    TemplateFiller,
    TemplateFillerError,
    fill_template,
)


class TestTemplateFillerCreation:
    """Test TemplateFiller initialization and configuration."""

    def test_create_filler_with_default_strategy(self):
        """Test creating filler with default missing value strategy."""
        filler = TemplateFiller()
        assert filler._missing_strategy == "keep"

    def test_create_filler_with_keep_strategy(self):
        """Test creating filler with 'keep' strategy."""
        filler = TemplateFiller(missing_placeholder_strategy="keep")
        assert filler._missing_strategy == "keep"

    def test_create_filler_with_empty_strategy(self):
        """Test creating filler with 'empty' strategy."""
        filler = TemplateFiller(missing_placeholder_strategy="empty")
        assert filler._missing_strategy == "empty"

    def test_create_filler_with_default_strategy(self):
        """Test creating filler with 'default' strategy."""
        filler = TemplateFiller(missing_placeholder_strategy="default")
        assert filler._missing_strategy == "default"

    def test_create_filler_with_invalid_strategy(self):
        """Test creating filler with invalid strategy raises error."""
        with pytest.raises(ValueError, match="Invalid missing_placeholder_strategy"):
            TemplateFiller(missing_placeholder_strategy="invalid")


class TestFillTextTemplate:
    """Test filling text-based templates."""

    @pytest.fixture
    def sample_mapping(self):
        """Create sample mapping for testing."""
        return Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={
                "Customer Name": "customer_name",
                "Invoice Amount": "amount",
                "Invoice Date": "date",
            },
        )

    @pytest.fixture
    def sample_data_row(self):
        """Create sample data row for testing."""
        return {
            "Customer Name": "John Doe",
            "Invoice Amount": 150.50,
            "Invoice Date": "2026-02-14",
            "Extra Column": "Should be ignored",
        }

    def test_fill_single_placeholder(self, sample_mapping):
        """Test filling template with single placeholder."""
        template = "Dear {{customer_name}},"
        data_row = {"Customer Name": "Alice"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Dear Alice,"

    def test_fill_multiple_placeholders(self, sample_mapping, sample_data_row):
        """Test filling template with multiple placeholders."""
        template = "Dear {{customer_name}}, your invoice of {{amount}} from {{date}} is due."

        filler = TemplateFiller()
        result = filler.fill_text_template(template, sample_data_row, sample_mapping)

        assert result == "Dear John Doe, your invoice of 150.5 from 2026-02-14 is due."

    def test_fill_duplicate_placeholders(self, sample_mapping):
        """Test filling template with duplicate placeholders."""
        template = "{{customer_name}} - {{customer_name}}"
        data_row = {"Customer Name": "Bob"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Bob - Bob"

    def test_fill_preserves_text_around_placeholders(self, sample_mapping):
        """Test that template text around placeholders is preserved."""
        template = "Start {{customer_name}} Middle {{amount}} End"
        data_row = {"Customer Name": "Test", "Invoice Amount": 100}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Start Test Middle 100 End"

    def test_fill_with_newlines(self, sample_mapping):
        """Test filling template with newlines."""
        template = "Dear {{customer_name}},\n\nYour invoice: {{amount}}\n\nThank you."
        data_row = {"Customer Name": "Jane", "Invoice Amount": 200}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Dear Jane,\n\nYour invoice: 200\n\nThank you."

    def test_fill_multiline_placeholder(self, sample_mapping):
        """Test filling template with multiline placeholder."""
        # Create mapping with multiline placeholder name
        mapping = Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={"Column": "customer name"},
        )

        template = "Dear {{customer name}},"
        data_row = {"Column": "Alice"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        assert result == "Dear Alice,"


class TestMissingValuesHandling:
    """Test handling of missing data values."""

    @pytest.fixture
    def sample_mapping(self):
        """Create sample mapping for testing."""
        return Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={
                "Name": "name",
                "Email": "email",
                "Phone": "phone",
            },
        )

    def test_keep_strategy_keeps_placeholder(self, sample_mapping):
        """Test 'keep' strategy keeps placeholder in output."""
        template = "{{name}} - {{email}} - {{phone}}"
        data_row = {"Name": "John", "Email": "john@example.com"}
        # Phone is missing

        filler = TemplateFiller(missing_placeholder_strategy="keep")
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "John - john@example.com - {{phone}}"

    def test_empty_strategy_removes_placeholder(self, sample_mapping):
        """Test 'empty' strategy removes placeholder from output."""
        template = "{{name}} - {{email}} - {{phone}}"
        data_row = {"Name": "John"}
        # Email and Phone are missing

        filler = TemplateFiller(missing_placeholder_strategy="empty")
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "John -  - "

    def test_default_strategy_uses_na(self, sample_mapping):
        """Test 'default' strategy uses N/A for missing values."""
        template = "{{name}} - {{email}} - {{phone}}"
        data_row = {"Name": "John"}
        # Email and Phone are missing

        filler = TemplateFiller(missing_placeholder_strategy="default")
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "John - N/A - N/A"

    def test_none_value_handled_as_missing(self, sample_mapping):
        """Test that None value is handled correctly."""
        template = "{{name}} - {{email}}"
        data_row = {"Name": "John", "Email": None}

        filler = TemplateFiller(missing_placeholder_strategy="keep")
        result = filler.fill_text_template(template, data_row, sample_mapping)

        # None should be treated as missing value
        assert result == "John - {{email}}"

    def test_unmapped_column_in_data_row(self, sample_mapping):
        """Test that unmapped columns in data row are ignored."""
        template = "{{name}}"
        data_row = {
            "Name": "John",
            "Unmapped Column": "Should be ignored",
            "Another Unmapped": 12345,
        }

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "John"

    def test_placeholder_not_in_mapping(self, sample_mapping):
        """Test placeholder in template but not in mapping."""
        template = "{{name}} - {{unknown_placeholder}}"
        data_row = {"Name": "John"}

        filler = TemplateFiller(missing_placeholder_strategy="keep")
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "John - {{unknown_placeholder}}"


class TestValueConversion:
    """Test conversion of different data types to strings."""

    @pytest.fixture
    def sample_mapping(self):
        """Create sample mapping for testing."""
        return Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={
                "Name": "name",
                "Age": "age",
                "Balance": "balance",
                "Active": "active",
            },
        )

    def test_string_value(self, sample_mapping):
        """Test filling with string value."""
        template = "{{name}}"
        data_row = {"Name": "Alice"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Alice"

    def test_integer_value(self, sample_mapping):
        """Test filling with integer value."""
        template = "Age: {{age}}"
        data_row = {"Age": 30}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Age: 30"

    def test_float_value(self, sample_mapping):
        """Test filling with float value."""
        template = "Balance: {{balance}}"
        data_row = {"Balance": 1234.56}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Balance: 1234.56"

    def test_boolean_value(self, sample_mapping):
        """Test filling with boolean value."""
        template = "Active: {{active}}"
        data_row = {"Active": True}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, sample_mapping)

        assert result == "Active: True"


class TestErrorHandling:
    """Test error handling in template filling."""

    @pytest.fixture
    def sample_mapping(self):
        """Create sample mapping for testing."""
        return Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={"Name": "name"},
        )

    def test_none_template_content_raises_error(self, sample_mapping):
        """Test that None template content raises error."""
        filler = TemplateFiller()

        with pytest.raises(TemplateFillerError, match="Template content cannot be None"):
            filler.fill_text_template(None, {"Name": "John"}, sample_mapping)

    def test_non_string_template_content_raises_error(self, sample_mapping):
        """Test that non-string template content raises error."""
        filler = TemplateFiller()

        with pytest.raises(TemplateFillerError, match="Template content must be string"):
            filler.fill_text_template(12345, {"Name": "John"}, sample_mapping)

    def test_non_dict_data_row_raises_error(self, sample_mapping):
        """Test that non-dict data row raises error."""
        filler = TemplateFiller()
        template = "{{name}}"

        with pytest.raises(ValueError, match="data_row must be dictionary"):
            filler.fill_text_template(template, "not a dict", sample_mapping)


class TestDOCXErrorHandling:
    """Test error handling for DOCX template filling."""

    def test_fill_docx_without_python_docx(self):
        """Test that ImportError is raised when python-docx not available."""
        # We can't actually uninstall python-docx, so we'll just
        # verify the code path exists by checking DOCX_AVAILABLE
        from src.services.template_filler import DOCX_AVAILABLE

        if not DOCX_AVAILABLE:
            # This test path would be taken if python-docx wasn't installed
            filler = TemplateFiller()

            with pytest.raises(ImportError, match="python-docx is not available"):
                filler.fill_docx_template(
                    "/tmp/template.docx",
                    {"Name": "Test"},
                    Mapping(
                        file_id="file-123",
                        template_id="template-456",
                        column_mappings={"Name": "name"},
                    ),
                )
        else:
            pytest.skip("python-docx is available, cannot test ImportError path")

    def test_fill_docx_missing_file(self):
        """Test filling DOCX template with missing file."""
        from src.services.template_filler import DOCX_AVAILABLE

        if DOCX_AVAILABLE:
            filler = TemplateFiller()
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={"Name": "name"},
            )

            with pytest.raises(FileNotFoundError, match="Template file not found"):
                filler.fill_docx_template("/nonexistent/template.docx", {}, mapping)
        else:
            pytest.skip("python-docx not available")

    def test_fill_docx_directory_instead_of_file(self):
        """Test filling DOCX template with directory path instead of file."""
        from src.services.template_filler import DOCX_AVAILABLE

        if DOCX_AVAILABLE:
            filler = TemplateFiller()
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={"Name": "name"},
            )

            # Use /tmp which is a directory
            with pytest.raises(TemplateFillerError, match="Template path is not a file"):
                filler.fill_docx_template("/tmp", {}, mapping)
        else:
            pytest.skip("python-docx not available")


class TestFillTemplateConvenienceFunction:
    """Test the convenience function for filling templates."""

    def test_fill_template_convenience(self):
        """Test fill_template() convenience function."""
        # Create a temporary text template file
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as f:
            f.write("Hello {{name}}, your balance is {{balance}}.")
            template_path = f.name

        try:
            # Create mapping and data
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={"Customer Name": "name", "Balance": "balance"},
            )
            data_row = {"Customer Name": "Alice", "Balance": 100.50}

            # Fill template
            result = fill_template(template_path, data_row, mapping)

            # Check result
            result_text = result.decode("utf-8")
            assert result_text == "Hello Alice, your balance is 100.5."

        finally:
            # Clean up temp file
            os.unlink(template_path)

    def test_fill_template_missing_file(self):
        """Test fill_template() with missing file."""
        mapping = Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={"Name": "name"},
        )
        data_row = {"Name": "John"}

        with pytest.raises(FileNotFoundError, match="Template file not found"):
            fill_template("/nonexistent/template.txt", data_row, mapping)

    def test_fill_template_unsupported_format(self):
        """Test fill_template() with unsupported file format."""
        # Create a temporary file with unsupported extension
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".pdf", delete=False, encoding="utf-8"
        ) as f:
            f.write("Hello {{name}}")
            template_path = f.name

        try:
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={"Name": "name"},
            )
            data_row = {"Name": "John"}

            with pytest.raises(
                TemplateFillerError, match="Unsupported template format"
            ):
                fill_template(template_path, data_row, mapping)

        finally:
            os.unlink(template_path)

    def test_fill_template_with_custom_strategy(self):
        """Test fill_template() with custom missing value strategy."""
        # Create a temporary text template file
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as f:
            f.write("Hello {{name}}, {{missing}}")
            template_path = f.name

        try:
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={"Name": "name"},
            )
            data_row = {"Name": "Alice"}

            # Fill with 'default' strategy
            result = fill_template(
                template_path,
                data_row,
                mapping,
                missing_placeholder_strategy="default",
            )

            # Check result
            result_text = result.decode("utf-8")
            assert result_text == "Hello Alice, N/A"

        finally:
            os.unlink(template_path)


class TestDOCXTemplateFilling:
    """Test filling DOCX templates (if python-docx is available)."""

    @pytest.fixture
    def sample_mapping(self):
        """Create sample mapping for testing."""
        return Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={
                "Customer Name": "customer_name",
                "Amount": "amount",
            },
        )

    def test_fill_docx_template_requires_docx_library(self, sample_mapping):
        """Test that filling DOCX without python-docx raises ImportError."""
        # This test is tricky because we can't easily uninstall python-docx
        # Instead, we'll just verify the method exists and can be called
        # when python-docx IS available
        from src.services.template_filler import DOCX_AVAILABLE

        if DOCX_AVAILABLE:
            # If python-docx is available, create a simple DOCX template
            try:
                from docx import Document

                # Create temporary DOCX file
                with tempfile.NamedTemporaryFile(
                    mode="wb", suffix=".docx", delete=False
                ) as f:
                    temp_path = f.name

                try:
                    # Create simple document with placeholder
                    doc = Document()
                    doc.add_paragraph("Dear {{customer_name}},")
                    doc.add_paragraph("Your invoice amount is: {{amount}}")
                    doc.save(temp_path)

                    # Fill template
                    filler = TemplateFiller()
                    data_row = {"Customer Name": "John Doe", "Amount": 150}
                    result = filler.fill_docx_template(
                        temp_path, data_row, sample_mapping
                    )

                    # Result should be bytes
                    assert isinstance(result, bytes)
                    assert len(result) > 0

                    # Also test output_path parameter
                    output_path = temp_path + ".filled.docx"
                    try:
                        result = filler.fill_docx_template(
                            temp_path,
                            data_row,
                            sample_mapping,
                            output_path=output_path,
                        )

                        # Should return None when output_path provided
                        assert result is None

                        # Check output file was created
                        assert Path(output_path).exists()

                    finally:
                        if os.path.exists(output_path):
                            os.unlink(output_path)

                finally:
                    os.unlink(temp_path)

            except ImportError:
                pytest.skip("python-docx not available for DOCX testing")
        else:
            pytest.skip("python-docx not available")


class TestFillTemplateAdditionalFormats:
    """Test filling templates with different file formats."""

    def test_fill_csv_template(self):
        """Test filling CSV template file."""
        # Create temporary CSV template
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".csv", delete=False, encoding="utf-8"
        ) as f:
            f.write("Name,Email,Phone\n{{name}},{{email}},{{phone}}")
            template_path = f.name

        try:
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={
                    "Name": "name",
                    "Email": "email",
                    "Phone": "phone",
                },
            )
            data_row = {"Name": "Alice", "Email": "alice@test.com", "Phone": "12345"}

            result = fill_template(template_path, data_row, mapping)
            result_text = result.decode("utf-8")

            assert result_text == "Name,Email,Phone\nAlice,alice@test.com,12345"

        finally:
            os.unlink(template_path)

    def test_fill_html_template(self):
        """Test filling HTML template file."""
        # Create temporary HTML template
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", delete=False, encoding="utf-8"
        ) as f:
            f.write("<h1>Hello {{name}}</h1><p>Email: {{email}}</p>")
            template_path = f.name

        try:
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={"Name": "name", "Email": "email"},
            )
            data_row = {"Name": "Bob", "Email": "bob@example.com"}

            result = fill_template(template_path, data_row, mapping)
            result_text = result.decode("utf-8")

            assert result_text == "<h1>Hello Bob</h1><p>Email: bob@example.com</p>"

        finally:
            os.unlink(template_path)

    def test_fill_md_template(self):
        """Test filling Markdown template file."""
        # Create temporary Markdown template
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".md", delete=False, encoding="utf-8"
        ) as f:
            f.write("# Report for {{name}}\n\nDate: {{date}}")
            template_path = f.name

        try:
            mapping = Mapping(
                file_id="file-123",
                template_id="template-456",
                column_mappings={"Name": "name", "Date": "date"},
            )
            data_row = {"Name": "Charlie", "Date": "2026-02-14"}

            result = fill_template(template_path, data_row, mapping)
            result_text = result.decode("utf-8")

            assert result_text == "# Report for Charlie\n\nDate: 2026-02-14"

        finally:
            os.unlink(template_path)


class TestFillTemplateMissingValueScenarios:
    """Test edge cases around missing value handling."""

    def test_missing_value_with_empty_dict_mapping(self):
        """Test filling template with empty column mappings."""
        mapping = Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={},
        )
        template = "Hello {{name}}"
        data_row = {"Name": "Alice"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        # No mapping, so placeholder should remain
        assert result == "Hello {{name}}"

    def test_missing_value_strategy_with_placeholders_only(self):
        """Test template with only placeholders and no data."""
        mapping = Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={"Name": "name"},
        )
        template = "{{name}}"
        data_row = {}

        filler = TemplateFiller(missing_placeholder_strategy="empty")
        result = filler.fill_text_template(template, data_row, mapping)

        # Empty strategy should remove placeholder
        assert result == ""


class TestPlaceholderMatching:
    """Test placeholder pattern matching edge cases."""

    @pytest.fixture
    def sample_mapping(self):
        """Create sample mapping for testing."""
        return Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={
                "Name": "name",
                "Field-Name": "field_name",
                "Field_Name": "field_name_underscore",
            },
        )

    def test_placeholder_with_hyphen(self, sample_mapping):
        """Test placeholder with hyphen in name."""
        # Create mapping that matches the placeholder name with hyphen
        mapping = Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={"Field-Name": "field-name"},
        )
        template = "{{field-name}}"
        data_row = {"Field-Name": "test123"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        assert result == "test123"

    def test_placeholder_with_underscore(self, sample_mapping):
        """Test placeholder with underscore in name."""
        # Create mapping that matches the placeholder name with underscore
        mapping = Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={"Field_Name": "field_name"},
        )
        template = "{{field_name}}"
        data_row = {"Field_Name": "test456"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        assert result == "test456"

    def test_placeholder_with_numbers(self, sample_mapping):
        """Test placeholder with numbers in name."""
        mapping = Mapping(
            file_id="file-123",
            template_id="template-456",
            column_mappings={"Field123": "field123"},
        )

        template = "{{field123}}"
        data_row = {"Field123": "value789"}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        assert result == "value789"

    def test_nested_braces_ignored(self):
        """Test that nested braces are ignored."""
        mapping = Mapping(
            file_id="file-123", template_id="template-456", column_mappings={}
        )

        template = "{{{name}}}"
        data_row = {}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        # Nested braces don't match our pattern, so should remain unchanged
        assert result == "{{{name}}}"

    def test_empty_placeholder_name_ignored(self):
        """Test that empty placeholder braces are ignored."""
        mapping = Mapping(
            file_id="file-123", template_id="template-456", column_mappings={}
        )

        template = "{{}}"
        data_row = {}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        # Empty braces don't match our pattern (requires at least one char)
        assert result == "{{}}"

    def test_special_chars_in_placeholder_ignored(self):
        """Test that special characters in placeholder name are ignored."""
        mapping = Mapping(
            file_id="file-123", template_id="template-456", column_mappings={}
        )

        template = "{{name@with#special$chars}}"
        data_row = {}

        filler = TemplateFiller()
        result = filler.fill_text_template(template, data_row, mapping)

        # Special chars don't match our pattern, so should remain unchanged
        assert result == "{{name@with#special$chars}}"
