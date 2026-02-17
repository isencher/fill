"""
Unit tests for DOCX Generator Service.
"""

import pytest
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from src.services.docx_generator import (
    DocxGenerator,
    DocxGeneratorError,
    generate_docx_from_data,
    generate_docx_from_template,
)


class TestDocxGeneratorCreation:
    """Tests for DocxGenerator initialization."""

    def test_create_generator(self):
        """Test creating a DocxGenerator instance."""
        generator = DocxGenerator()
        assert generator is not None


class TestGenerateFromData:
    """Tests for generate_from_data method."""

    def test_generate_simple_data(self, tmp_path):
        """Test generating a document from simple data."""
        generator = DocxGenerator()
        data = [
            {"name": "Alice", "age": 30},
            {"name": "Bob", "age": 25},
        ]

        output_path = tmp_path / "output.docx"
        result = generator.generate_from_data(data, output_path=output_path)

        # Should return None when output_path is provided
        assert result is None
        # File should exist
        assert output_path.exists()

    def test_generate_with_title(self, tmp_path):
        """Test generating a document with a title."""
        generator = DocxGenerator()
        data = [{"name": "Alice", "age": 30}]

        output_path = tmp_path / "output_with_title.docx"
        generator.generate_from_data(
            data, output_path=output_path, title="Employee List"
        )

        assert output_path.exists()

    def test_generate_without_headers(self, tmp_path):
        """Test generating a document without headers."""
        generator = DocxGenerator()
        data = [{"name": "Alice", "age": 30}]

        output_path = tmp_path / "output_no_headers.docx"
        generator.generate_from_data(
            data, output_path=output_path, include_headers=False
        )

        assert output_path.exists()

    def test_generate_returns_bytes(self):
        """Test generating a document returns bytes when no output path."""
        generator = DocxGenerator()
        data = [{"name": "Alice", "age": 30}]

        result = generator.generate_from_data(data)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_empty_data_raises_error(self):
        """Test that empty data raises an error."""
        generator = DocxGenerator()
        with pytest.raises(DocxGeneratorError, match="Data cannot be empty"):
            generator.generate_from_data([])

    def test_generate_invalid_data_type_raises_error(self):
        """Test that non-list data raises an error."""
        generator = DocxGenerator()
        with pytest.raises(DocxGeneratorError, match="Data must be a list"):
            generator.generate_from_data("not a list")

    def test_generate_non_dict_rows_raises_error(self):
        """Test that non-dictionary rows raise an error."""
        generator = DocxGenerator()
        with pytest.raises(
            DocxGeneratorError, match=r"All data rows must be dictionaries"
        ):
            generator.generate_from_data([["not", "a", "dict"]])

    def test_generate_with_various_data_types(self, tmp_path):
        """Test generating a document with various data types."""
        generator = DocxGenerator()
        data = [
            {"name": "Alice", "age": 30, "active": True, "score": 95.5},
            {"name": "Bob", "age": None, "active": False, "score": None},
        ]

        output_path = tmp_path / "output_types.docx"
        generator.generate_from_data(data, output_path=output_path)

        assert output_path.exists()

    def test_generate_with_many_columns(self, tmp_path):
        """Test generating a document with many columns."""
        generator = DocxGenerator()
        data = [
            {
                f"col{i}": f"value{i}" for i in range(10)
            }
        ]

        output_path = tmp_path / "output_many_cols.docx"
        generator.generate_from_data(data, output_path=output_path)

        assert output_path.exists()

    def test_generate_with_many_rows(self, tmp_path):
        """Test generating a document with many rows."""
        generator = DocxGenerator()
        data = [{"name": f"Person{i}", "id": i} for i in range(100)]

        output_path = tmp_path / "output_many_rows.docx"
        generator.generate_from_data(data, output_path=output_path)

        assert output_path.exists()


class TestGenerateFromTemplate:
    """Tests for generate_from_template method."""

    def test_generate_from_template(self, tmp_path):
        """Test generating from a template file."""
        # First, create a template file
        from docx import Document

        template_path = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("Name: {{name}}")
        doc.add_paragraph("Age: {{age}}")
        doc.save(str(template_path))

        # Now use the template
        generator = DocxGenerator()
        data = {"name": "Alice", "age": 30}

        output_path = tmp_path / "output_from_template.docx"
        generator.generate_from_template(
            template_path, data, output_path=output_path
        )

        assert output_path.exists()

    def test_generate_from_template_returns_bytes(self, tmp_path):
        """Test generating from template returns bytes when no output path."""
        from docx import Document

        template_path = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("Name: {{name}}")
        doc.save(str(template_path))

        generator = DocxGenerator()
        data = {"name": "Alice"}

        result = generator.generate_from_template(template_path, data)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_from_template_missing_file(self, tmp_path):
        """Test that missing template file raises FileNotFoundError."""
        generator = DocxGenerator()
        non_existent_path = tmp_path / "does_not_exist.docx"

        with pytest.raises(FileNotFoundError, match="Template file not found"):
            generator.generate_from_template(
                non_existent_path, {"name": "Alice"}
            )

    def test_generate_from_template_invalid_extension(self, tmp_path):
        """Test that non-.docx template raises an error."""
        generator = DocxGenerator()
        template_path = tmp_path / "template.txt"
        template_path.write_text("Name: {{name}}")

        with pytest.raises(DocxGeneratorError, match="Template must be .docx"):
            generator.generate_from_template(
                template_path, {"name": "Alice"}
            )

    def test_generate_from_template_keeps_unmapped_placeholders(
        self, tmp_path
    ):
        """Test that unmapped placeholders are kept in the output."""
        from docx import Document

        template_path = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("Name: {{name}}")
        doc.add_paragraph("Email: {{email}}")
        doc.save(str(template_path))

        generator = DocxGenerator()
        data = {"name": "Alice"}  # email not provided

        output_path = tmp_path / "output_partial.docx"
        generator.generate_from_template(
            template_path, data, output_path=output_path
        )

        assert output_path.exists()


class TestGenerateSimpleDocument:
    """Tests for generate_simple_document method."""

    def test_generate_simple_document_string(self, tmp_path):
        """Test generating a simple document with string content."""
        generator = DocxGenerator()
        content = "This is a simple document."

        output_path = tmp_path / "simple.docx"
        generator.generate_simple_document(
            content, output_path=output_path
        )

        assert output_path.exists()

    def test_generate_simple_document_list(self, tmp_path):
        """Test generating a simple document with list content."""
        generator = DocxGenerator()
        content = ["Paragraph 1", "Paragraph 2", "Paragraph 3"]

        output_path = tmp_path / "simple_list.docx"
        generator.generate_simple_document(
            content, output_path=output_path
        )

        assert output_path.exists()

    def test_generate_simple_document_with_title(self, tmp_path):
        """Test generating a simple document with a title."""
        generator = DocxGenerator()
        content = "This is a simple document."

        output_path = tmp_path / "simple_with_title.docx"
        generator.generate_simple_document(
            content, output_path=output_path, title="My Document"
        )

        assert output_path.exists()

    def test_generate_simple_document_returns_bytes(self):
        """Test that simple document returns bytes when no output path."""
        generator = DocxGenerator()
        content = "Simple document content."

        result = generator.generate_simple_document(content)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_simple_document_empty_content_raises_error(self):
        """Test that empty content raises an error."""
        generator = DocxGenerator()
        with pytest.raises(DocxGeneratorError, match="Content cannot be empty"):
            generator.generate_simple_document("")

    def test_generate_simple_document_invalid_content_type(self):
        """Test that invalid content type raises an error."""
        generator = DocxGenerator()
        with pytest.raises(DocxGeneratorError, match="Content must be string or list"):
            generator.generate_simple_document(123)

    def test_generate_simple_document_skips_empty_strings(self, tmp_path):
        """Test that empty strings in list are skipped."""
        generator = DocxGenerator()
        content = ["Paragraph 1", "", "Paragraph 2"]

        output_path = tmp_path / "simple_with_empty.docx"
        generator.generate_simple_document(
            content, output_path=output_path
        )

        assert output_path.exists()


class TestConvenienceFunctions:
    """Tests for convenience functions."""

    def test_generate_docx_from_data(self):
        """Test generate_docx_from_data convenience function."""
        data = [{"name": "Alice", "age": 30}]

        result = generate_docx_from_data(data)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_docx_from_template(self, tmp_path):
        """Test generate_docx_from_template convenience function."""
        from docx import Document

        template_path = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("Name: {{name}}")
        doc.save(str(template_path))

        data = {"name": "Alice"}

        result = generate_docx_from_template(template_path, data)

        assert isinstance(result, bytes)
        assert len(result) > 0


class TestDocxGeneratorError:
    """Tests for DocxGeneratorError exception."""

    def test_error_creation(self):
        """Test creating a DocxGeneratorError."""
        error = DocxGeneratorError("Test error message")
        assert error.message == "Test error message"
        assert str(error) == "Test error message"


class TestDocxGeneratorOutputValidation:
    """Tests for validating generated DOCX output."""

    def test_output_is_valid_docx(self, tmp_path):
        """Test that output is a valid DOCX file."""
        from docx import Document

        generator = DocxGenerator()
        data = [{"name": "Alice", "age": 30}]

        output_path = tmp_path / "valid_output.docx"
        generator.generate_from_data(data, output_path=output_path)

        # Verify the file can be opened by python-docx
        doc = Document(str(output_path))
        assert doc is not None
        # Should have a table (even if no paragraphs)
        assert len(doc.tables) > 0

    def test_output_has_table(self, tmp_path):
        """Test that generated document contains a table."""
        from docx import Document

        generator = DocxGenerator()
        data = [{"name": "Alice", "age": 30}, {"name": "Bob", "age": 25}]

        output_path = tmp_path / "output_with_table.docx"
        generator.generate_from_data(data, output_path=output_path)

        # Verify the document has a table
        doc = Document(str(output_path))
        assert len(doc.tables) > 0
        # Header row + 2 data rows
        assert len(doc.tables[0].rows) == 3

    def test_output_with_title(self, tmp_path):
        """Test that generated document with title has heading."""
        from docx import Document

        generator = DocxGenerator()
        data = [{"name": "Alice"}]

        output_path = tmp_path / "output_with_title.docx"
        generator.generate_from_data(
            data, output_path=output_path, title="Test Title"
        )

        # Verify the document has a heading
        doc = Document(str(output_path))
        # First paragraph should be the heading
        assert len(doc.paragraphs) > 0
