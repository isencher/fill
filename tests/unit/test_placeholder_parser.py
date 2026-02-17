"""
Unit tests for Placeholder Parser Service.

Tests cover placeholder extraction from text and files,
validation, edge cases, and error handling.
"""

import pytest
from pathlib import Path

from src.services.placeholder_parser import (
    PlaceholderParser,
    get_placeholder_parser,
)


class TestPlaceholderParserCreation:
    """Test PlaceholderParser initialization."""

    def test_new_parser(self):
        """Test creating a new parser."""
        parser = PlaceholderParser()
        assert parser is not None

    def test_get_placeholder_parser_singleton(self):
        """Test that get_placeholder_parser returns singleton."""
        parser1 = get_placeholder_parser()
        parser2 = get_placeholder_parser()

        assert parser1 is parser2


class TestExtractFromText:
    """Test extract_from_text operation."""

    def test_extract_single_placeholder(self):
        """Test extracting a single placeholder."""
        parser = PlaceholderParser()
        content = "Hello {{name}}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["name"]

    def test_extract_multiple_placeholders(self):
        """Test extracting multiple placeholders."""
        parser = PlaceholderParser()
        content = "Dear {{customer}}, your order {{order_id}} is ready."

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["customer", "order_id"]

    def test_extract_no_placeholders(self):
        """Test extracting from text with no placeholders."""
        parser = PlaceholderParser()
        content = "Hello world, no placeholders here."

        placeholders = parser.extract_from_text(content)

        assert placeholders == []

    def test_extract_with_duplicates(self):
        """Test that duplicates are preserved."""
        parser = PlaceholderParser()
        content = "{{name}} and {{name}} again"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["name", "name"]

    def test_extract_with_numbers(self):
        """Test extracting placeholders with numbers."""
        parser = PlaceholderParser()
        content = "Item {{item_1}} and {{item_2}}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["item_1", "item_2"]

    def test_extract_with_hyphens(self):
        """Test extracting placeholders with hyphens."""
        parser = PlaceholderParser()
        content = "Field {{field-name}} and {{another-field}}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["field-name", "another-field"]

    def test_extract_with_underscores(self):
        """Test extracting placeholders with underscores."""
        parser = PlaceholderParser()
        content = "{{first_name}} {{last_name}}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["first_name", "last_name"]

    def test_extract_with_spaces_in_placeholder(self):
        """Test extracting placeholders with internal spaces."""
        parser = PlaceholderParser()
        content = "{{field name}} and {{another field}}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["field name", "another field"]

    def test_extract_whitespace_trimmed(self):
        """Test that leading/trailing whitespace is trimmed."""
        parser = PlaceholderParser()
        content = "{{  name  }} and {{  value  }}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["name", "value"]

    def test_extract_multiline_placeholder(self):
        """Test extracting multiline placeholders."""
        parser = PlaceholderParser()
        content = """Hello {{
            name
        }}"""

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["name"]

    def test_extract_mixed_characters(self):
        """Test extracting placeholders with mixed valid characters."""
        parser = PlaceholderParser()
        content = "{{field_1-2}} and {{field_A-B}}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["field_1-2", "field_A-B"]

    def test_extract_preserves_order(self):
        """Test that extraction preserves placeholder order."""
        parser = PlaceholderParser()
        content = "{{first}} {{second}} {{third}}"

        placeholders = parser.extract_from_text(content)

        assert placeholders == ["first", "second", "third"]

    def test_extract_from_empty_string(self):
        """Test extracting from empty string."""
        parser = PlaceholderParser()
        content = ""

        placeholders = parser.extract_from_text(content)

        assert placeholders == []

    def test_extract_with_special_chars_ignored(self):
        """Test that special chars in placeholders are ignored."""
        parser = PlaceholderParser()
        # Only alphanumeric, underscore, hyphen, space are valid
        # The regex won't match {{field$}}
        content = "{{field1}} and {{field$}} and {{field2}}"

        placeholders = parser.extract_from_text(content)

        # field$ won't match because $ is not valid
        assert "field1" in placeholders
        assert "field2" in placeholders


class TestExtractFromTextErrors:
    """Test extract_from_text error handling."""

    def test_none_content_raises_error(self):
        """Test that None content raises ValueError."""
        parser = PlaceholderParser()
        with pytest.raises(ValueError, match="cannot be None"):
            parser.extract_from_text(None)

    def test_non_string_content_raises_error(self):
        """Test that non-string content raises ValueError."""
        parser = PlaceholderParser()
        with pytest.raises(ValueError, match="must be string"):
            parser.extract_from_text(123)


class TestExtractUniqueFromText:
    """Test extract_unique_from_text operation."""

    def test_extract_unique_removes_duplicates(self):
        """Test that duplicates are removed."""
        parser = PlaceholderParser()
        content = "{{name}} and {{name}} and {{value}}"

        placeholders = parser.extract_unique_from_text(content)

        assert placeholders == ["name", "value"]

    def test_extract_unique_preserves_order(self):
        """Test that unique extraction preserves first occurrence order."""
        parser = PlaceholderParser()
        content = "{{c}} {{b}} {{a}} {{b}} {{c}} {{a}}"

        placeholders = parser.extract_unique_from_text(content)

        assert placeholders == ["c", "b", "a"]

    def test_extract_unique_no_duplicates(self):
        """Test unique with no duplicates."""
        parser = PlaceholderParser()
        content = "{{a}} {{b}} {{c}}"

        placeholders = parser.extract_unique_from_text(content)

        assert placeholders == ["a", "b", "c"]

    def test_extract_unique_empty(self):
        """Test unique extraction from empty content."""
        parser = PlaceholderParser()
        content = ""

        placeholders = parser.extract_unique_from_text(content)

        assert placeholders == []


class TestExtractFromFile:
    """Test extract_from_file operation."""

    def test_extract_from_text_file(self, tmp_path: Path):
        """Test extracting from a text file."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.txt"
        test_file.write_text("Hello {{name}}, your order {{order_id}} is ready.")

        placeholders = parser.extract_from_file(test_file)

        assert placeholders == ["name", "order_id"]

    def test_extract_from_markdown_file(self, tmp_path: Path):
        """Test extracting from a markdown file."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.md"
        test_file.write_text("# {{title}}\n\nAuthor: {{author}}")

        placeholders = parser.extract_from_file(test_file)

        assert placeholders == ["title", "author"]

    def test_extract_from_html_file(self, tmp_path: Path):
        """Test extracting from an HTML file."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.html"
        test_file.write_text("<h1>{{title}}</h1><p>{{content}}</p>")

        placeholders = parser.extract_from_file(test_file)

        assert placeholders == ["title", "content"]

    def test_extract_from_csv_file(self, tmp_path: Path):
        """Test extracting from a CSV file."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.csv"
        test_file.write_text("{{field1}},{{field2}},{{field3}}")

        placeholders = parser.extract_from_file(test_file)

        assert placeholders == ["field1", "field2", "field3"]

    def test_extract_from_file_with_string_path(self, tmp_path: Path):
        """Test extracting with string path instead of Path object."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.txt"
        test_file.write_text("{{field1}} {{field2}}")

        placeholders = parser.extract_from_file(str(test_file))

        assert placeholders == ["field1", "field2"]

    def test_extract_from_file_with_multiline(self, tmp_path: Path):
        """Test extracting from file with multiline content."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.txt"
        test_file.write_text("Line 1 {{field1}}\nLine 2 {{field2}}\nLine 3")

        placeholders = parser.extract_from_file(test_file)

        assert placeholders == ["field1", "field2"]


class TestExtractFromFileErrors:
    """Test extract_from_file error handling."""

    def test_nonexistent_file_raises_error(self, tmp_path: Path):
        """Test that nonexistent file raises ValueError."""
        parser = PlaceholderParser()
        with pytest.raises(ValueError, match="not found"):
            parser.extract_from_file(tmp_path / "nonexistent.txt")

    def test_directory_raises_error(self, tmp_path: Path):
        """Test that directory path raises ValueError."""
        parser = PlaceholderParser()
        with pytest.raises(ValueError, match="not a file"):
            parser.extract_from_file(tmp_path)

    def test_extract_unique_from_file(self, tmp_path: Path):
        """Test extract_unique_from_file removes duplicates."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.txt"
        test_file.write_text("{{name}} {{name}} {{value}}")

        placeholders = parser.extract_unique_from_file(test_file)

        assert placeholders == ["name", "value"]

    def test_unicode_decode_error_in_text_file(self, tmp_path: Path):
        """Test that non-UTF8 file raises ValueError."""
        parser = PlaceholderParser()
        test_file = tmp_path / "test.txt"
        # Write bytes that are not valid UTF-8
        test_file.write_bytes(b"\xff\xfe Invalid UTF-8")

        with pytest.raises(ValueError, match="encoding issue"):
            parser.extract_from_file(test_file)


class TestValidatePlaceholder:
    """Test validate_placeholder operation."""

    def test_validate_valid_placeholder(self):
        """Test validating a valid placeholder."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("field_name") is True

    def test_validate_with_numbers(self):
        """Test validating placeholder with numbers."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("field123") is True

    def test_validate_with_underscores(self):
        """Test validating placeholder with underscores."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("field_name_test") is True

    def test_validate_with_hyphens(self):
        """Test validating placeholder with hyphens."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("field-name") is True

    def test_validate_with_spaces(self):
        """Test validating placeholder with spaces."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("field name") is True

    def test_validate_with_mixed_valid_chars(self):
        """Test validating placeholder with mixed valid characters."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("field_name-123 Test") is True

    def test_validate_empty_string(self):
        """Test validating empty string."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("") is False

    def test_validate_none(self):
        """Test validating None."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder(None) is False

    def test_validate_non_string(self):
        """Test validating non-string value."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder(123) is False

    def test_validate_with_braces(self):
        """Test that placeholder with braces is invalid."""
        parser = PlaceholderParser()
        # Braces are not valid in field names themselves
        assert parser.validate_placeholder("{{name}}") is False

    def test_validate_with_special_chars(self):
        """Test that special characters are invalid."""
        parser = PlaceholderParser()
        assert parser.validate_placeholder("field$") is False
        assert parser.validate_placeholder("field@") is False
        assert parser.validate_placeholder("field#") is False


class TestExtractFromDocx:
    """Test DOCX file extraction (if python-docx is available)."""

    def test_extract_from_docx_raises_error_without_docx(self, tmp_path: Path):
        """Test that DOCX extraction raises ImportError without python-docx."""
        # Create a parser with docx disabled
        parser = PlaceholderParser()
        parser._docx_available = False

        test_file = tmp_path / "test.docx"
        test_file.write_text("fake docx content")

        with pytest.raises(ImportError, match="python-docx is required"):
            parser.extract_from_file(test_file)

    def test_extract_from_docx_invalid_file(self, tmp_path: Path):
        """Test that invalid DOCX file raises ValueError."""
        # Skip if python-docx is not available
        parser = PlaceholderParser()
        if not parser._docx_available:
            pytest.skip("python-docx not available")

        # Create a fake docx file (not actually valid)
        test_file = tmp_path / "test.docx"
        test_file.write_bytes(b"not a valid docx file")

        with pytest.raises(ValueError, match="Failed to open"):
            parser.extract_from_file(test_file)

    def test_extract_from_docx_paragraphs(self, tmp_path: Path):
        """Test extracting placeholders from DOCX paragraphs."""
        parser = PlaceholderParser()
        if not parser._docx_available:
            pytest.skip("python-docx not available")

        # Create a real DOCX file with placeholders
        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Hello {{name}}, your order {{order_id}} is ready.")
            doc.add_paragraph("Total: {{total_amount}}")

            test_file = tmp_path / "test.docx"
            doc.save(str(test_file))

            placeholders = parser.extract_from_file(test_file)

            assert "name" in placeholders
            assert "order_id" in placeholders
            assert "total_amount" in placeholders
        except ImportError:
            pytest.skip("python-docx not available")

    def test_extract_from_docx_tables(self, tmp_path: Path):
        """Test extracting placeholders from DOCX tables."""
        parser = PlaceholderParser()
        if not parser._docx_available:
            pytest.skip("python-docx not available")

        try:
            from docx import Document

            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "{{field1}}"
            table.rows[0].cells[1].text = "{{field2}}"
            table.rows[1].cells[0].text = "Data 1"
            table.rows[1].cells[1].text = "Data 2"

            test_file = tmp_path / "test.docx"
            doc.save(str(test_file))

            placeholders = parser.extract_from_file(test_file)

            assert "field1" in placeholders
            assert "field2" in placeholders
        except ImportError:
            pytest.skip("python-docx not available")

    def test_extract_from_docx_mixed(self, tmp_path: Path):
        """Test extracting placeholders from DOCX with both paragraphs and tables."""
        parser = PlaceholderParser()
        if not parser._docx_available:
            pytest.skip("python-docx not available")

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Dear {{customer}}")
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "{{invoice_number}}"
            table.rows[0].cells[1].text = "{{date}}"
            doc.add_paragraph("Total: {{amount}}")

            test_file = tmp_path / "test.docx"
            doc.save(str(test_file))

            placeholders = parser.extract_from_file(test_file)

            assert "customer" in placeholders
            assert "invoice_number" in placeholders
            assert "date" in placeholders
            assert "amount" in placeholders
        except ImportError:
            pytest.skip("python-docx not available")

    def test_extract_unique_from_docx(self, tmp_path: Path):
        """Test extract_unique_from_file removes duplicates from DOCX."""
        parser = PlaceholderParser()
        if not parser._docx_available:
            pytest.skip("python-docx not available")

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("{{name}} and {{name}} and {{value}}")

            test_file = tmp_path / "test.docx"
            doc.save(str(test_file))

            placeholders = parser.extract_unique_from_file(test_file)

            assert placeholders == ["name", "value"]
        except ImportError:
            pytest.skip("python-docx not available")
